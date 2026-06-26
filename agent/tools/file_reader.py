"""
============================================================
File Reader 工具 — 读取本地文件内容
============================================================

让 Agent 具备读取本地文件的能力，支持：
- 文本文件（.txt, .md, .py, .json, .csv 等）
- PDF 文档（.pdf）— 依赖 PyMuPDF (fitz)
- Word 文档（.docx）— 依赖 python-docx
- 范围读取（start_line / end_line 分页）— 大文件的核心兜底机制
- 文件结构概览（Python 函数/类定义、Markdown 标题等）
- 限制最大读取行数/字符数，防止超大文件撑爆上下文
- 安全的文件存在性检查

安全机制：
    - 扩展名白名单过滤（30+ 种文本格式）
    - 最大行数限制（默认 200 行）
    - 文件大小保护（超过 10MB 拒绝读取）
    - UTF-8 编码强制
    - 路径安全检查（ToolContext.assert_file_path_allowed）

为什么支持范围读取（兜底机制）：
    - LLM 上下文窗口有限，一次性读完整个大文件会浪费 token 预算
    - Agent 可以先读头部了解文件结构，再按需读取关键部分
    - 分页读取让大文件分析成为可能——即使被截断也有办法继续
    - 当 max_result_chars / max_result_tokens 触发截断后，LLM 可以通过
      调整 start_line 精确读取下一段内容，避免重复读取已看过的行
"""

from pathlib import Path
from agent.tools.base import BaseTool, ToolResult, ToolError, ErrorCode, require_import


# ============================================================
# 允许读取的文件扩展名白名单
# 防止 Agent 尝试读取二进制文件（如图片、可执行文件等）
# 无扩展名的文件（如 .env, Makefile）也允许
# ============================================================
ALLOWED_EXTENSIONS = {
    ".txt", ".md", ".py", ".json", ".csv", ".xml", ".html",
    ".css", ".js", ".ts", ".java", ".c", ".cpp", ".h", ".go",
    ".rs", ".rb", ".php", ".sh", ".bat", ".yaml", ".yml",
    ".toml", ".ini", ".cfg", ".conf", ".log", ".sql",
    ".r", ".swift", ".kt", ".scala", ".lua", ".env",
    ".gitignore", ".dockerfile",
    ".pdf", ".docx",
}


class FileReaderTool(BaseTool):
    """
    文件读取工具 — 让 Agent 能读取本地文本文件。

    支持范围读取（start_line/end_line），方便 Agent 分页查看大文件。
    即使被 BaseTool 的 max_result_chars / max_result_tokens 截断，
    LLM 也可以通过新的范围参数精确请求下一段内容。

    工作原理：
    1. 接收文件路径参数 + 可选的范围参数（start_line / end_line / max_lines）
    2. 检查文件是否存在、是否为文件（非目录）
    3. 检查扩展名是否在白名单中
    4. 根据文件类型选择读取方式（文本/PDF/Word）
    5. 计算有效范围，读取指定行区间
    6. 对 Python/Markdown 等文件自动提取结构概览
    7. 返回格式化的文件内容 + 范围元数据 + 分页提示
    """

    @property
    def name(self) -> str:
        return "read_file"

    @property
    def timeout(self) -> float:
        return 30.0

    @property
    def quota_limit(self) -> int:
        """
        文件读取工具的单会话配额上限（token 数）。

        文件内容可能很大，设为 5000 tokens（约 7500 字符）。
        修改为 0 可取消限额，仅统计使用量。
        """
        return 5000

    @property
    def description(self) -> str:
        return (
            "读取指定路径的文件内容。"
            "支持 .txt, .md, .py, .json, .csv, .xml, .yaml 等文本文件，"
            "以及 .pdf 文档和 .docx Word 文档。"
            "可用于分析代码、查看配置、读取数据文件、检查日志、阅读文档等。"
            "文件过大时可通过 start_line/end_line 参数分页读取，避免截断。"
        )

    @property
    def parameters(self) -> dict:
        return {
            "type": "object",
            "properties": {
                "file_path": {
                    "type": "string",
                    "description": "要读取的文件路径（支持绝对路径和相对路径）",
                },
                "max_lines": {
                    "type": "integer",
                    "description": (
                        "最大读取行数，默认 200。当不指定 end_line 时，"
                        "从 start_line 开始读取 max_lines 行。"
                    ),
                    "default": 200,
                },
                "start_line": {
                    "type": "integer",
                    "description": (
                        "起始行号（从1开始），默认从第1行开始。"
                        "对于大文件，先不指定此参数读取头部了解文件结构，"
                        "然后通过 start_line/end_line 分页查看关键部分。"
                    ),
                    "default": 1,
                },
                "end_line": {
                    "type": "integer",
                    "description": (
                        "结束行号（包含此行）。与 start_line 配合实现精确范围读取。"
                        "例如 start_line=100, end_line=200 表示读取第 100 到 200 行。"
                        "不设置此参数时使用 max_lines 控制读取行数。"
                    ),
                },
            },
            "required": ["file_path"],
        }

    def before_execute(self, kwargs: dict, context) -> dict | None:
        """
        读取前检查路径权限（由 before_execute 钩子自动调用）。
        确保 Agent 不会读到不允许访问的路径。
        """
        file_path = kwargs.get("file_path", "")
        if file_path and context:
            context.assert_file_path_allowed(file_path, operation="read")
        return None

    def get_dynamic_parameters(self, current_params: dict) -> dict:
        """
        根据文件路径动态调整 max_lines 的推荐值。

        不同类型的文件推荐不同的行数，帮助 LLM 给出更合理的默认值。
        """
        file_path = current_params.get("file_path", "")
        if not file_path:
            return {}
        suffix = Path(file_path).suffix.lower()
        suggestions = {
            ".py": {"max_lines": 200, "desc": "Python 源码推荐读取 200 行"},
            ".json": {"max_lines": 500, "desc": "JSON 文件可能很紧凑，可适当增加行数"},
            ".csv": {"max_lines": 100, "desc": "CSV 文件行数多但每行短，注意截断"},
            ".log": {"max_lines": 100, "desc": "日志文件通常很大，建议限制行数"},
            ".md": {"max_lines": 200, "desc": "Markdown 文件适中"},
            ".pdf": {"max_lines": 100, "desc": "PDF 文本提取后行数可能较多"},
        }
        if suffix in suggestions:
            info = suggestions[suffix]
            return {
                "max_lines": {
                    "type": "integer",
                    "description": f"最大读取行数，默认 {info['max_lines']}。{info['desc']}",
                    "default": info["max_lines"],
                }
            }
        return {}

    def execute(
        self,
        file_path: str,
        max_lines: int = 200,
        start_line: int = 1,
        end_line: int | None = None,
    ) -> ToolResult:
        """
        执行文件读取。

        执行流程：
        1. 路径解析与安全检查
        2. 检查白名单和文件大小
        3. 计算有效读取范围（end_line 优先于 max_lines）
        4. 按文件类型分流读取（文本/PDF/Word）
        5. 返回格式化的文件内容 + 元数据（含范围信息 + 分页提示）

        Args:
            file_path: 文件路径字符串
            max_lines: 最大读取行数（当 end_line 不设置时生效）
            start_line: 起始行号（从1开始），默认从第1行开始
            end_line: 结束行号（包含），与 start_line 配合实现精确范围读取

        Returns:
            ToolResult: 包含文件内容、范围信息、文件结构的元数据
        """
        try:
            path = Path(file_path).expanduser().resolve()

            # 检查文件是否存在
            if not path.exists():
                return ToolResult(
                    success=False, result="",
                    error=f"文件不存在: {path}",
                )

            # 检查是否为文件（排除目录）
            if not path.is_file():
                return ToolResult(
                    success=False, result="",
                    error=f"路径不是文件: {path}",
                )

            # 检查扩展名白名单（无扩展名的文件如 .env 也允许）
            suffix = path.suffix.lower()
            if suffix and suffix not in ALLOWED_EXTENSIONS:
                return ToolResult(
                    success=False, result="",
                    error=(
                        f"不支持的文件类型: {suffix}。"
                        f"支持的类型: {', '.join(sorted(ALLOWED_EXTENSIONS)[:12])}..."
                    ),
                )

            # 检查文件大小（10MB 保护）
            file_size = path.stat().st_size
            if file_size > 10 * 1024 * 1024:
                return ToolResult(
                    success=False, result="",
                    error=f"文件过大: {file_size / 1024 / 1024:.1f}MB，超过 10MB 限制",
                )

            # ============================================================
            # 计算有效读取范围
            # 优先级: end_line > max_lines
            # end_line 是精确的结束行号（包含），max_lines 是相对 start_line 的偏移量
            # ============================================================
            if end_line is not None and end_line >= start_line:
                # 用户指定了精确范围：读取 start_line ~ end_line
                effective_max = end_line - start_line + 1
            else:
                # 用户未指定 end_line，从 start_line 开始读 max_lines 行
                effective_max = max_lines

            # 按文件类型分流读取
            if suffix == ".pdf":
                content, read_meta = self._read_pdf(path, start_line, effective_max)
            elif suffix == ".docx":
                content, read_meta = self._read_docx(path, start_line, effective_max)
            else:
                content, read_meta = self._read_text(path, start_line, effective_max)

            # 构建 continuation 指引（供 BaseTool 截断时使用）
            # 如果文件还有更多内容未显示，声明如何分页继续读取
            has_more = read_meta.get("has_more", False)
            continuation = None
            if has_more:
                next_start = read_meta.get("actual_end_line", start_line) + 1
                continuation = {
                    "tool": "read_file",
                    "params": {
                        "file_path": str(path),
                        "start_line": next_start,
                        "max_lines": effective_max,
                    },
                    "hint": (
                        f"需要继续读取时，请调用 read_file 并设置 "
                        f"start_line={next_start}, max_lines={effective_max}"
                    ),
                }

            # 合并元数据并返回
            return ToolResult(
                success=True,
                result=content,
                continuation=continuation,
                metadata={
                    "path": str(path),
                    "file_size": file_size,
                    "file_type": suffix if suffix else "text",
                    "total_lines": read_meta.get("total_lines", 0),
                    "start_line": start_line,
                    "end_line": read_meta.get("actual_end_line", 0),
                    "displayed_lines": read_meta.get("displayed_lines", 0),
                    "has_more": has_more,
                },
            )

        except PermissionError:
            return ToolResult(
                success=False, result="",
                error=f"权限不足，无法读取文件: {file_path}",
            )
        except Exception as e:
            return ToolResult(
                success=False, result="",
                error=f"读取文件失败: {type(e).__name__}: {e}",
            )

    # ============================================================
    # 文件结构分析（辅助方法）
    # 主要是为了方便模型阅读又不占用大量上下文篇幅和token
    # ============================================================

    def _get_file_structure(self, path: Path, lines: list[str]) -> list[str]:
        """
        分析文件结构，提取关键结构标记。

        对不同文件类型执行不同的分析策略：
        - .py: 提取顶层函数定义(def)和类定义(class)
        - .md: 提取所有标题(# ## 等)
        - .json: 提取顶层键名

        为什么需要这个分析：
        在范围读取模式下，LLM 先读文件头部获得结构概览，
        然后可以针对性地跳转到感兴趣的部分查看详细内容。

        Args:
            path: 文件路径（用于判断文件类型）
            lines: 文件的所有行（已经读入内存，扫描几乎无开销）

        Returns:
            list[str]: 结构标记行列表（最多返回 30 条，防止结果过大）
        """
        if not lines:
            return []

        suffix = path.suffix.lower()
        markers = []
        seen = set()

        if suffix == ".py":
            # Python 文件：提取顶层的 def 和 class（不含缩进的内容）
            for line in lines:
                # 只匹配无缩进的顶层定义（模块级的函数和类）
                if not line or line[0] in (" ", "\t"):
                    continue
                stripped = line.strip()
                if stripped.startswith("def "):
                    key = stripped[:stripped.index("(")] if "(" in stripped else stripped
                    if key not in seen:
                        # 提取函数名：def xxx(...) → xxx
                        func_name = stripped[4:stripped.index("(")] if "(" in stripped else stripped[4:]
                        markers.append(f"  def {func_name}()")
                        seen.add(key)
                elif stripped.startswith("class "):
                    # 提取类名：class Xxx(...) → Xxx
                    colon_idx = stripped.index(":") if ":" in stripped else len(stripped)
                    paren_idx = stripped.index("(") if "(" in stripped else colon_idx
                    class_name = stripped[6:min(paren_idx, colon_idx)].strip()
                    key = f"class {class_name}"
                    if key not in seen:
                        markers.append(f"  class {class_name}")
                        seen.add(key)

        elif suffix == ".md":
            # Markdown 文件：提取各级标题行
            for line in lines:
                stripped = line.strip()
                if stripped.startswith("#") and not stripped.startswith("# "):
                    # 跳过"# "这种非标题的用法
                    pass
                if stripped.startswith("#"):
                    markers.append(f"  {stripped}")

        elif suffix == ".json":
            # JSON 文件：提取顶层键
            import json as json_mod
            try:
                full_text = "\n".join(lines)
                data = json_mod.loads(full_text)
                if isinstance(data, dict):
                    for key in list(data.keys())[:30]:
                        markers.append(f"  键: {key}")
                elif isinstance(data, list):
                    markers.append(f"  列表，共 {len(data)} 项")
            except json_mod.JSONDecodeError:
                pass

        return markers[:30]  # 最多返回 30 条，防止结果过大

    # ============================================================
    # 专用读取方法
    # ============================================================

    def _read_text(self, path: Path, start_line: int, max_lines: int) -> tuple[str, dict]:
        """
        读取纯文本文件（UTF-8 编码），支持范围读取。

        这是最常用的读取方式，所有非 PDF/Word 文件都走此方法。
        范围读取是核心兜底机制：
        - 当文件被 max_result_chars 截断时，显示"总行数"和"当前范围"
        - LLM 可以据此决定下一次读取的行区间
        - 避免跨上下文传递时重复读取已看过的内容

        Args:
            path: 文件路径
            start_line: 起始行号（1-based），从第几行开始读
            max_lines: 要读取的最大行数

        Returns:
            tuple[str, dict]: (格式化内容字符串, 元数据字典)
                元数据包含 total_lines, actual_end_line, displayed_lines, has_more
        """
        content = path.read_text(encoding="utf-8")
        all_lines = content.splitlines()
        total_lines = len(all_lines)

        # 计算行范围：将 1-based 的 start_line 转为 0-based 索引
        start_idx = max(0, start_line - 1)
        end_idx = min(total_lines, start_idx + max_lines)

        selected = all_lines[start_idx:end_idx]
        actual_end_line = min(total_lines, start_line + max_lines - 1)

        # ============================================================
        # 文件结构分析（仅对超过 50 行的文件执行）
        # 让 LLM 在头部就看到文件大纲，按需跳转阅读
        # ============================================================
        structure_hint = ""
        if total_lines > 50:
            markers = self._get_file_structure(path, all_lines)
            if markers:
                structure_hint = (
                    "\n📄 文件结构概览:\n"
                    + "\n".join(markers)
                    + "\n\n"
                )

        # 构建文件头信息：包含范围、大小、结构
        header = (
            f"📁 文件: {path}\n"
            f"📏 大小: {path.stat().st_size:,} 字节 | 总行数: {total_lines}\n"
        )
        if total_lines > max_lines or start_line > 1:
            header += (
                f"🎯 当前范围: 第 {start_line}~{actual_end_line} 行 "
                f"(显示 {len(selected)} 行, 共 {total_lines} 行)\n"
            )
        header += f"{'=' * 60}\n"

        # 内容部分
        content_text = "\n".join(selected) if selected else "(空)\n"

        # ============================================================
        # 分页提示（核心兜底机制）
        # 当文件还有更多内容未显示时，给出明确的下次读取提示。
        # 这样即使 BaseTool 截断了结果，LLM 也知道下一步怎么读。
        # ============================================================
        has_more = actual_end_line < total_lines
        pagination_hint = ""
        if has_more:
            next_start = actual_end_line + 1
            pagination_hint = (
                f"\n{'=' * 60}\n"
                f"💡 提示: 文件共 {total_lines} 行，当前仅显示了第 {start_line}~{actual_end_line} 行。\n"
                f"   需要继续读取时，请调用 read_file 并设置 "
                f"start_line={next_start}, max_lines={max_lines}\n"
            )

        result = header + structure_hint + content_text + pagination_hint

        metadata = {
            "total_lines": total_lines,
            "actual_end_line": actual_end_line,
            "displayed_lines": len(selected),
            "has_more": has_more,
        }

        return result, metadata

    @require_import("fitz", "PyMuPDF")
    def _read_pdf(self, path: Path, start_line: int, max_lines: int) -> tuple[str, dict]:
        """
        读取 PDF 文件 — 使用 PyMuPDF (fitz) 提取文本，支持范围读取。

        PDF 的"行"是提取后文本的行（非原始 PDF 的行布局）。
        支持分页提示，方便 Agent 逐段阅读长文档。

        Args:
            path: 文件路径
            start_line: 起始行号（1-based，针对提取后的文本行）
            max_lines: 要读取的最大行数

        Returns:
            tuple[str, dict]: (格式化内容字符串, 元数据字典)
        """
        import fitz

        doc = fitz.open(path)
        total_pages = len(doc)

        # 提取所有文本到行列表，同时记录分页位置
        all_lines = []

        for page_num in range(total_pages):
            page = doc[page_num]
            text = page.get_text()
            page_lines = text.splitlines()
            page_lines = [l for l in page_lines if l.strip()]

            if page_lines:
                all_lines.append(f"--- 第 {page_num + 1} 页 ---")
                all_lines.extend(page_lines)

        doc.close()
        total_lines = len(all_lines)

        # 计算行范围（1-based → 0-based）
        start_idx = max(0, start_line - 1)
        end_idx = min(total_lines, start_idx + max_lines)

        selected = all_lines[start_idx:end_idx]
        actual_end_line = min(total_lines, start_line + max_lines - 1)

        # 构建文件头
        header = (
            f"📁 文件: {path}\n"
            f"📏 大小: {path.stat().st_size:,} 字节 | PDF 页数: {total_pages}\n"
        )
        if total_lines > max_lines or start_line > 1:
            header += (
                f"🎯 当前范围: 第 {start_line}~{actual_end_line} 行 "
                f"(显示 {len(selected)} 行, 共 {total_lines} 行)\n"
            )
        header += f"{'=' * 60}\n"

        content_text = "\n".join(selected) if selected else "(空)\n"

        # 分页提示
        has_more = actual_end_line < total_lines
        pagination_hint = ""
        if has_more:
            next_start = actual_end_line + 1
            pagination_hint = (
                f"\n{'=' * 60}\n"
                f"💡 提示: PDF 文本共 {total_lines} 行，当前仅显示第 {start_line}~{actual_end_line} 行。\n"
                f"   需要继续读取时，请调用 read_file 并设置 "
                f"start_line={next_start}, max_lines={max_lines}\n"
            )

        result = header + content_text + pagination_hint

        metadata = {
            "total_lines": total_lines,
            "actual_end_line": actual_end_line,
            "displayed_lines": len(selected),
            "has_more": has_more,
            "total_pages": total_pages,
        }

        return result, metadata

    @require_import("docx", "python-docx")
    def _read_docx(self, path: Path, start_line: int, max_lines: int) -> tuple[str, dict]:
        """
        读取 Word (.docx) 文件 — 使用 python-docx 提取文本，支持范围读取。

        先提取段落文本，如果不够 max_lines，再提取表格内容。
        所有内容组装成行列表后按范围截取。

        Args:
            path: 文件路径
            start_line: 起始行号（1-based）
            max_lines: 要读取的最大行数

        Returns:
            tuple[str, dict]: (格式化内容字符串, 元数据字典)
        """
        from docx import Document

        doc = Document(path)

        # 提取所有内容到行列表（段落 + 表格）
        all_lines = []

        # 提取段落文本
        for para in doc.paragraphs:
            if para.text.strip():
                all_lines.append(para.text)

        # 提取表格内容（段落后追加）
        for table in doc.tables:
            all_lines.append("")
            all_lines.append("--- 表格 ---")
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                all_lines.append(" | ".join(cells))

        total_lines = len(all_lines)

        # 计算行范围（1-based → 0-based）
        start_idx = max(0, start_line - 1)
        end_idx = min(total_lines, start_idx + max_lines)

        selected = all_lines[start_idx:end_idx]
        actual_end_line = min(total_lines, start_line + max_lines - 1)

        # 构建文件头
        header = (
            f"📁 文件: {path}\n"
            f"📏 大小: {path.stat().st_size:,} 字节 | 段落数: {len(doc.paragraphs)}\n"
        )
        if total_lines > max_lines or start_line > 1:
            header += (
                f"🎯 当前范围: 第 {start_line}~{actual_end_line} 行 "
                f"(显示 {len(selected)} 行, 共 {total_lines} 行)\n"
            )
        header += f"{'=' * 60}\n"

        content_text = "\n".join(selected) if selected else "(空)\n"

        # 分页提示
        has_more = actual_end_line < total_lines
        pagination_hint = ""
        if has_more:
            next_start = actual_end_line + 1
            pagination_hint = (
                f"\n{'=' * 60}\n"
                f"💡 提示: 文档共 {total_lines} 行，当前仅显示第 {start_line}~{actual_end_line} 行。\n"
                f"   需要继续读取时，请调用 read_file 并设置 "
                f"start_line={next_start}, max_lines={max_lines}\n"
            )

        result = header + content_text + pagination_hint

        metadata = {
            "total_lines": total_lines,
            "actual_end_line": actual_end_line,
            "displayed_lines": len(selected),
            "has_more": has_more,
        }

        return result, metadata
